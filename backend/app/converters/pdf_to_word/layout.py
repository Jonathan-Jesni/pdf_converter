import pdfplumber
from docx import Document
import io
import os

LINE_Y_THRESHOLD = 2
MIN_TEXT_CHARS = 30
TAB_WIDTH = 40
COLUMN_GAP_THRESHOLD = 100


def is_meaningful_text(words, min_chars=MIN_TEXT_CHARS):
    total_chars = sum(len(w["text"].strip()) for w in words)
    return total_chars >= min_chars


def add_full_width_image(doc, image_buffer):
    section = doc.sections[-1]
    max_width = section.page_width - section.left_margin - section.right_margin
    doc.add_picture(image_buffer, width=max_width)


def group_words_into_lines(words):
    lines = []
    current = []

    for w in sorted(words, key=lambda x: (x["top"], x["x0"])):
        if not current:
            current.append(w)
            continue

        if abs(w["top"] - current[-1]["top"]) <= LINE_Y_THRESHOLD:
            current.append(w)
        else:
            lines.append(current)
            current = [w]

    if current:
        lines.append(current)

    return lines


def split_into_columns(words):
    words = sorted(words, key=lambda w: w["x0"])
    columns = []
    current = [words[0]]

    for w in words[1:]:
        if abs(w["x0"] - current[-1]["x0"]) > COLUMN_GAP_THRESHOLD:
            columns.append(current)
            current = [w]
        else:
            current.append(w)

    columns.append(current)
    return columns


def render_column(doc, words):
    lines = group_words_into_lines(words)
    left_margin = min(w["x0"] for w in words)

    for line in lines:
        text = " ".join(w["text"] for w in line).strip()
        if not text:
            continue

        x0 = line[0]["x0"]
        indent = int((x0 - left_margin) / TAB_WIDTH)
        indent = max(0, indent)

        doc.add_paragraph("\t" * indent + text)

def extract_page_images(page):
    images = []

    for img in page.images:
        try:
            bbox = (
                img["x0"],
                img["top"],
                img["x1"],
                img["bottom"]
            )
            cropped = page.crop(bbox).to_image(resolution=300).original
            buf = io.BytesIO()
            cropped.save(buf, format="PNG")
            buf.seek(0)
            images.append(buf)
        except Exception:
            continue

    return images


def pdf_to_word_layout(input_pdf_path, output_docx_path):
    doc = Document()
    os.makedirs(os.path.dirname(output_docx_path), exist_ok=True)

    with pdfplumber.open(input_pdf_path) as pdf:
        for page in pdf.pages:
            words = page.extract_words(use_text_flow=True)

            # Scanned / image-only fallback
            if not words or not is_meaningful_text(words):
                img = page.to_image(resolution=300).original
                buf = io.BytesIO()
                img.save(buf, format="PNG")
                buf.seek(0)
                add_full_width_image(doc, buf)
                doc.add_page_break()
                continue

            columns = split_into_columns(words)

            for col_words in columns:
                render_column(doc, col_words)

            # ADD embedded images AFTER text
            images = extract_page_images(page)
            for img_buf in images:
                add_full_width_image(doc, img_buf)

            doc.add_page_break()


    doc.save(output_docx_path)


if __name__ == "__main__":
    pdf_to_word_layout(
        "app/storage/uploads/input.pdf",
        "app/storage/outputs/output_layout.docx"
    )
    print("Layout mode conversion finished")

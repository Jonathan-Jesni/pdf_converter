import pdfplumber
from docx import Document
from .layout import pdf_to_word_layout, render_layout
from backend.app.core.analysis.build_profile import build_page_profile
import io
import os


LINE_Y_THRESHOLD = 3
PARAGRAPH_Y_GAP = 12
HEADING_SCALE = 1.25
MIN_TEXT_CHARS = 30
COLUMN_GAP_THRESHOLD = 50
ROW_Y_THRESHOLD = 10


def add_full_width_image(doc, image_buffer):
    section = doc.sections[-1]
    max_width = section.page_width - section.left_margin - section.right_margin
    doc.add_picture(image_buffer, width=max_width)


def is_meaningful_text(words, min_chars=MIN_TEXT_CHARS):
    total_chars = sum(len(w["text"].strip()) for w in words)
    return total_chars >= min_chars


def group_words_into_lines(words):
    lines = []
    current = []

    for w in sorted(words, key=lambda x: (x["top"], x["x0"])):
        if not current:
            current.append(w)
            continue

        if abs(w["top"] - current[-1]["top"]) <= LINE_Y_THRESHOLD:
            current.append(w)
        else:
            lines.append(current)
            current = [w]

    if current:
        lines.append(current)

    return lines


def split_into_columns(words):
    words = sorted(words, key=lambda w: w["x0"])
    columns = []
    current = [words[0]]

    for w in words[1:]:
        if abs(w["x0"] - current[-1]["x0"]) > COLUMN_GAP_THRESHOLD:
            columns.append(current)
            current = [w]
        else:
            current.append(w)

    columns.append(current)
    return columns


def extract_lines(words):
    lines = group_words_into_lines(words)
    result = []
    for line in lines:
        text = " ".join(w["text"] for w in line).strip()
        top = line[0]["top"]
        if text:
            result.append((top, text))
    return result


def pair_form_rows(left_lines, right_lines):
    pairs = []
    used = set()

    for l_top, l_text in left_lines:
        best = None
        best_diff = None

        for i, (r_top, r_text) in enumerate(right_lines):
            if i in used:
                continue
            diff = abs(l_top - r_top)
            if diff <= ROW_Y_THRESHOLD and (best_diff is None or diff < best_diff):
                best = (i, r_text)
                best_diff = diff

        if best:
            idx, r_text = best
            used.add(idx)
            pairs.append((l_text, r_text))
        else:
            pairs.append((l_text, ""))

    return pairs


def pdf_to_word_no_ocr(
    input_pdf_path,
    output_docx_path,
    mode="semantic",
    report_path=None,
    pages=None
):
    if mode == "layout":
        pdf_to_word_layout(input_pdf_path, output_docx_path)
        return

    doc = Document()
    output_dir = os.path.dirname(output_docx_path)
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)

    decision_log = []

    with pdfplumber.open(input_pdf_path) as pdf:

        # -------- COLLECT PAGES --------
        page_items = [
            (idx, page)
            for idx, page in enumerate(pdf.pages, start=1)
            if pages is None or idx in pages
        ]

        # -------- PASS 1: ANALYSIS --------
        profiles = []
        for idx, page in page_items:
            words = page.extract_words(use_text_flow=True)

            profile = build_page_profile(
                page_number=idx,
                words=words,
                images=[]
            )

            profiles.append(profile)

        # -------- PASS 2: RENDER --------
        for profile in profiles:
            page = pdf.pages[profile.page_number - 1]
            words = profile.words
            page_mode = profile.detected_mode

            decision_log.append({
                "page": profile.page_number,
                "mode": page_mode,
                "reason": profile.reason
            })

            # ---- Image-only fallback ----
            if not words or not is_meaningful_text(words):
                img = page.to_image(resolution=300).original
                buf = io.BytesIO()
                img.save(buf, format="PNG")
                buf.seek(0)
                add_full_width_image(doc, buf)
                doc.add_page_break()
                continue

            # ---- Table rendering ----
            if page_mode == "table":
                table = doc.add_table(
                    rows=len(profile.table_cells),
                    cols=len(profile.table_cells[0])
                )

                for i, row in enumerate(profile.table_cells):
                    for j, cell in enumerate(row):
                        table.cell(i, j).text = cell

                doc.add_page_break()
                continue


            # ---- Layout rendering ----
            if page_mode == "layout":
                render_layout(profile, doc)
                doc.add_page_break()
                continue

            # ---- Form rendering ----
            if page_mode == "form":
                columns = split_into_columns(words)
                if len(columns) == 2:
                    left_lines = extract_lines(columns[0])
                    right_lines = extract_lines(columns[1])
                    pairs = pair_form_rows(left_lines, right_lines)

                    table = doc.add_table(rows=len(pairs), cols=2)
                    for i, (l, r) in enumerate(pairs):
                        table.cell(i, 0).text = l
                        table.cell(i, 1).text = r

                    doc.add_page_break()
                    continue

            # ---- Structured semantic rendering ----

            # Render headings
            for heading in getattr(profile, "headings", []):
                doc.add_heading(heading, level=1)

            # Render lists
            for lst in getattr(profile, "lists", []):
                for item in lst:
                    doc.add_paragraph(item, style="List Bullet")

            # Render paragraphs
            for text in getattr(profile, "paragraphs", []):
                if text:
                    doc.add_paragraph(text)

            doc.add_page_break()

    if report_path and decision_log:
        import json
        with open(report_path, "w") as f:
            json.dump(decision_log, f, indent=2)

    doc.save(output_docx_path)
